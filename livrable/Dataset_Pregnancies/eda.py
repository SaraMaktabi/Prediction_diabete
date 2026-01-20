import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from docx import Document
from docx.shared import Inches

# ------------------------
# 1. Charger le dataset
# ------------------------
df = pd.read_csv(r"diabetes.csv")

# ------------------------
# 2. Prétraitement
# ------------------------
# Colonnes où 0 n'est pas possible médicalement
cols_with_zero_invalid = ["Glucose", "BloodPressure", "SkinThickness", "Insulin", "BMI"]

# Remplacer 0 par NaN
df[cols_with_zero_invalid] = df[cols_with_zero_invalid].replace(0, pd.NA)

# Imputer les NaN par la moyenne de la colonne
df[cols_with_zero_invalid] = df[cols_with_zero_invalid].fillna(df[cols_with_zero_invalid].mean())

# ------------------------
# 3. Créer le document Word
# ------------------------
doc = Document()
doc.add_heading("EDA - Dataset Pima Diabetes", 0)

# ------------------------
# 4. Statistiques descriptives
# ------------------------
doc.add_heading("Statistiques descriptives", level=1)
desc = df.describe().transpose()
doc.add_paragraph(desc.to_string())

# ------------------------
# 5. Distribution de la variable cible
# ------------------------
doc.add_heading("Distribution de la variable cible (Outcome)", level=1)
target_counts = df["Outcome"].value_counts()
doc.add_paragraph(target_counts.to_string())

# ------------------------
# 6. Histogrammes et boxplots pour toutes les variables
# ------------------------
for col in df.columns.drop("Outcome"):
    # Histogramme
    plt.figure(figsize=(5,3))
    sns.histplot(df[col], kde=True, bins=20)
    plt.title(f"Histogramme - {col}")
    buf = BytesIO()
    plt.savefig(buf, format='png')
    plt.close()
    doc.add_heading(f"Histogramme - {col}", level=2)
    doc.add_picture(buf, width=Inches(5))

    # Boxplot
    plt.figure(figsize=(5,3))
    sns.boxplot(x=df[col])
    plt.title(f"Boxplot - {col}")
    buf = BytesIO()
    plt.savefig(buf, format='png')
    plt.close()
    doc.add_heading(f"Boxplot - {col}", level=2)
    doc.add_picture(buf, width=Inches(5))

# ------------------------
# 7. Corrélation et heatmap
# ------------------------
doc.add_heading("Matrice de corrélation - Heatmap", level=1)
numeric_cols = df.select_dtypes(include=['float64', 'int64']).columns
corr = df[numeric_cols].corr()

plt.figure(figsize=(10,8))
sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm", cbar=True)
plt.title("Heatmap des corrélations")
buf = BytesIO()
plt.savefig(buf, format='png')
plt.close()
doc.add_picture(buf, width=Inches(6))

# ------------------------
# 8. Statistiques par classe Outcome
# ------------------------
doc.add_heading("Statistiques par classe Outcome", level=1)
grouped = df.groupby("Outcome").mean()
doc.add_paragraph(grouped.to_string())

# ------------------------
# 9. Sauvegarde du document
# ------------------------
output_doc = r"EDA_dataset_pregnancies.docx"
doc.save(output_doc)
print(f"EDA complet sauvegardé dans : {output_doc}")
