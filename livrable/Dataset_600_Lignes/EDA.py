import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import os

# ===============================
# 1. Chargement des données
# ===============================
df = pd.read_csv("diabetes_data_upload.csv")

# ===============================
# 2. Création du document Word
# ===============================
doc = Document()
doc.add_heading("Analyse Exploratoire des Données (EDA)", level=1)

# Dossier pour images
IMG_DIR = "eda_images"
os.makedirs(IMG_DIR, exist_ok=True)

# ===============================
# 3. Histogrammes pour chaque feature
# ===============================
doc.add_heading("Histogrammes des variables numériques", level=2)

numeric_cols = df.select_dtypes(include=["int64", "float64"]).columns

for col in numeric_cols:
    plt.figure()
    sns.histplot(df[col], bins=20, kde=True)
    plt.title(f"Distribution de {col}")
    plt.xlabel(col)
    plt.ylabel("Fréquence")

    img_path = f"{IMG_DIR}/hist_{col}.png"
    plt.savefig(img_path)
    plt.close()

    doc.add_heading(f"Histogramme de {col}", level=3)
    doc.add_picture(img_path, width=Inches(5))

# ===============================
# 4. Boxplots (détection des outliers)
# ===============================
doc.add_heading("Boxplots pour la détection des valeurs aberrantes", level=2)

for col in numeric_cols:
    plt.figure()
    sns.boxplot(x=df[col])
    plt.title(f"Boxplot de {col}")

    img_path = f"{IMG_DIR}/box_{col}.png"
    plt.savefig(img_path)
    plt.close()

    doc.add_heading(f"Boxplot de {col}", level=3)
    doc.add_picture(img_path, width=Inches(5))

# ===============================
# 5. Distribution de la variable cible
# ===============================
doc.add_heading("Distribution de la variable cible", level=2)

plt.figure()
sns.countplot(x="class", data=df)
plt.title("Distribution de la variable cible")

img_path = f"{IMG_DIR}/class_distribution.png"
plt.savefig(img_path)
plt.close()

doc.add_picture(img_path, width=Inches(5))

# ===============================
# 6. Variables catégorielles vs classe
# ===============================
doc.add_heading("Relation entre les symptômes et la classe", level=2)

categorical_cols = df.select_dtypes(include="object").columns.drop("class")

for col in categorical_cols:
    plt.figure()
    sns.countplot(x=col, hue="class", data=df)
    plt.title(f"{col} selon la classe")
    plt.xticks(rotation=30)

    img_path = f"{IMG_DIR}/cat_{col.replace(' ', '_')}.png"
    plt.savefig(img_path)
    plt.close()

    doc.add_heading(f"{col} vs classe", level=3)
    doc.add_picture(img_path, width=Inches(5))

# ===============================
# 7. Heatmap de corrélation détaillée
# ===============================
doc.add_heading("Matrice de corrélation détaillée", level=2)

df_encoded = df.copy()

for col in df_encoded.columns:
    if df_encoded[col].dtype == "object":
        df_encoded[col] = df_encoded[col].map({
            "Yes": 1, "No": 0,
            "Male": 1, "Female": 0,
            "Positive": 1, "Negative": 0
        })

plt.figure(figsize=(14, 12))
sns.heatmap(
    df_encoded.corr(),
    annot=True,
    fmt=".2f",
    cmap="coolwarm",
    linewidths=0.5
)
plt.title("Heatmap de corrélation (détaillée)")

img_path = f"{IMG_DIR}/correlation_heatmap_detailed.png"
plt.savefig(img_path)
plt.close()

doc.add_picture(img_path, width=Inches(6))

# ===============================
# 8. Sauvegarde du document
# ===============================
doc.save("EDA_Dataset_600.docx")

print("✔ Document Word généré : EDA_Complete_Diabetes.docx")
