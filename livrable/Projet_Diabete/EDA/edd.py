import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import os

# ------------------------
# 1. Charger le dataset
# ------------------------
df = pd.read_csv("..\Dataset\diabetes_data_upload.csv") # remplacer par ton chemin

# Créer un document Word
doc = Document()
doc.add_heading('Analyse Exploratoire des Données (EDA) - Diabetes', 0)

# ------------------------
# 2. Informations générales
# ------------------------
doc.add_heading('1. Aperçu du dataset', level=1)
doc.add_paragraph(str(df.head()))
doc.add_paragraph(str(df.info()))

# ------------------------
# 3. Statistiques descriptives
# ------------------------
doc.add_heading('2. Statistiques descriptives', level=1)
doc.add_paragraph("=== Statistiques Age ===")
doc.add_paragraph(str(df["Age"].describe()))

categorical_cols = df.columns.drop("Age")
doc.add_paragraph("=== Distribution des variables catégorielles ===")
for col in categorical_cols:
    doc.add_paragraph(f"{col} :\n{df[col].value_counts()}")

# ------------------------
# 4. Transformation binaire
# ------------------------
df_bin = df.copy()
df_bin["class"] = df_bin["class"].map({"Positive":1, "Negative":0})
for col in categorical_cols:
    if col != "class":
        df_bin[col] = df_bin[col].map({"Yes":1, "No":0})

# ------------------------
# 5. Corrélation
# ------------------------
doc.add_heading('3. Corrélation', level=1)
doc.add_paragraph("=== Matrice complète ===")
doc.add_paragraph(str(df_bin.corr().round(2)))

doc.add_paragraph("=== Corrélation avec la variable cible class ===")
doc.add_paragraph(str(df_bin.corr()["class"].sort_values(ascending=False).round(2)))

# ------------------------
# 6. Graphiques
# ------------------------

# Fonction pour sauvegarder et ajouter les images dans Word
def add_plot_to_doc(fig, filename, doc, width=Inches(5)):
    fig.tight_layout()
    fig.savefig(filename)
    plt.close(fig)
    doc.add_picture(filename, width=width)
    os.remove(filename)

# Histogramme Age
fig = plt.figure(figsize=(6,4))
plt.hist(df["Age"], bins=10, color='skyblue', edgecolor='black')
plt.title("Distribution de l'âge")
plt.xlabel("Age")
plt.ylabel("Nombre de patients")
add_plot_to_doc(fig, "hist_age.png", doc)

# Heatmap corrélation
fig = plt.figure(figsize=(12,10))
sns.heatmap(df_bin.corr(), annot=True, cmap="coolwarm")
plt.title("Corrélation entre variables binaires et class")
add_plot_to_doc(fig, "heatmap_corr.png", doc, width=Inches(6))

# Boxplot Age vs class
fig = plt.figure(figsize=(6,4))
sns.boxplot(x="class", y="Age", data=df_bin)
plt.title("Age selon class (0=Negative, 1=Positive)")
add_plot_to_doc(fig, "boxplot_age.png", doc)

# ------------------------
# 7. Sauvegarde du document Word
# ------------------------
output_file = "EDA_Diabetes_2db.docx"
doc.save(output_file)
print(f"EDA complet sauvegardé dans : {output_file}")
