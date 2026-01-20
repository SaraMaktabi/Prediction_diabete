import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import os

# Charger le dataset
df = pd.read_csv("..\Dataset\Dataset_before_cleaning.csv")

# Créer un document Word
doc = Document()
doc.add_heading('Analyse Exploratoire des Données (EDA) - Diabetes', 0)

# ------------------------
# 1. Informations générales
# ------------------------
doc.add_heading('1. Informations générales', level=1)
doc.add_paragraph(str(df.info()))
doc.add_paragraph(str(df.describe()))

# ------------------------
# 2. Distribution de la variable cible
# ------------------------
doc.add_heading('2. Distribution de la variable cible', level=1)
diabetes_counts = df["Diabetes"].value_counts()
doc.add_paragraph(str(diabetes_counts))

plt.figure(figsize=(6,4))
diabetes_counts.plot(kind='bar', color=['skyblue','salmon'])
plt.title("Distribution de la variable Diabetes")
plt.xlabel("Diabetes")
plt.ylabel("Nombre de patients")
plt.tight_layout()
plt.savefig("diabetes_distribution.png")
plt.close()

doc.add_picture("diabetes_distribution.png", width=Inches(5))

# ------------------------
# 3. Corrélation des variables numériques
# ------------------------
doc.add_heading('3. Corrélation des variables numériques', level=1)
corr = df.corr(numeric_only=True)
doc.add_paragraph(str(corr))

plt.figure(figsize=(12,10))
sns.heatmap(corr, annot=True, fmt=".2f", cmap="coolwarm")
plt.title("Matrice de corrélation")
plt.tight_layout()
plt.savefig("correlation_matrix.png")
plt.close()
doc.add_picture("correlation_matrix.png", width=Inches(6))

# ------------------------
# 4. Boxplots variables clés vs Diabetes
# ------------------------
doc.add_heading('4. Boxplots des variables clés vs Diabetes', level=1)
numerical_vars = ["BMI", "Waist_Circumference", "Fasting_Blood_Glucose", "Blood_Pressure_Systolic", "Blood_Pressure_Diastolic"]
for var in numerical_vars:
    plt.figure(figsize=(6,4))
    sns.boxplot(x="Diabetes", y=var, data=df)
    plt.title(f"{var} selon le statut diabétique")
    plt.tight_layout()
    filename = f"boxplot_{var}.png"
    plt.savefig(filename)
    plt.close()
    doc.add_paragraph(f"Boxplot : {var} selon Diabetes")
    doc.add_picture(filename, width=Inches(5))

# ------------------------
# 5. Variables catégorielles vs Diabetes
# ------------------------
doc.add_heading('5. Variables catégorielles vs Diabetes', level=1)
categorical_vars = ["Sex", "Ethnicity", "Physical_Activity_Level", "Alcohol_Consumption", "Smoking_Status", "Family_History_of_Diabetes", "Previous_Gestational_Diabetes"]
for var in categorical_vars:
    crosstab = pd.crosstab(df[var], df["Diabetes"], normalize="index") * 100
    doc.add_paragraph(f"Tableau croisé {var} vs Diabetes (%) :\n{crosstab.round(2)}")

# ------------------------
# Sauvegarde du document Word
# ------------------------
output_file = "EDA_Diabetes.docx"
doc.save(output_file)
print(f"EDA complet sauvegardé dans : {output_file}")

# Nettoyage images temporaires
for f in os.listdir():
    if f.startswith("boxplot_") or f == "diabetes_distribution.png" or f == "correlation_matrix.png":
        os.remove(f)
