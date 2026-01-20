import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
from docx import Document
from docx.shared import Inches
import os


df = pd.read_csv("Dataset\Dataset_with_Diabetes.csv")
# ===============================
# 1. Création du document Word
# ===============================
doc = Document()
doc.add_heading("EDA - Dataset Clinique Diabetes", level=1)

IMG_DIR = "eda_images_clinical"
os.makedirs(IMG_DIR, exist_ok=True)

# ===============================
# 2. Histogrammes pour chaque variable numérique
# ===============================
numeric_cols = df.select_dtypes(include=["int64", "float64"]).columns

for col in numeric_cols:
    plt.figure()
    sns.histplot(df[col], bins=20, kde=True)
    plt.title(f"Distribution de {col}")
    plt.xlabel(col)
    plt.ylabel("Fréquence")
    
    img_path = f"{IMG_DIR}/hist_{col}.png"
    plt.savefig(img_path)
    plt.close()
    
    doc.add_heading(f"Histogramme de {col}", level=2)
    doc.add_picture(img_path, width=Inches(5))

# ===============================
# 3. Boxplots pour détecter outliers
# ===============================
for col in numeric_cols:
    plt.figure()
    sns.boxplot(x=df[col])
    plt.title(f"Boxplot de {col}")
    
    img_path = f"{IMG_DIR}/box_{col}.png"
    plt.savefig(img_path)
    plt.close()
    
    doc.add_heading(f"Boxplot de {col}", level=2)
    doc.add_picture(img_path, width=Inches(5))

# ===============================
# 4. Heatmap de corrélation détaillée
# ===============================
df_encoded = df.copy()
for col in df_encoded.select_dtypes(include="object").columns:
    df_encoded[col] = df_encoded[col].astype('category').cat.codes

plt.figure(figsize=(14,12))
sns.heatmap(df_encoded.corr(), annot=True, fmt=".2f", cmap="coolwarm", linewidths=0.5)
plt.title("Heatmap de corrélation détaillée")

img_path = f"{IMG_DIR}/correlation_heatmap.png"
plt.savefig(img_path)
plt.close()

doc.add_heading("Matrice de corrélation", level=2)
doc.add_picture(img_path, width=Inches(6))

# ===============================
# 5. Distribution cible
# ===============================
plt.figure()
sns.countplot(x='Diabetes', data=df)
plt.title("Distribution de la variable cible")
img_path = f"{IMG_DIR}/diabetes_distribution.png"
plt.savefig(img_path)
plt.close()
doc.add_heading("Distribution de la variable cible", level=2)
doc.add_picture(img_path, width=Inches(5))

# ===============================
# 6. Sauvegarde Word
# ===============================
doc.save("EDA_Dataset_10000.docx")
print("✔ Document Word généré : EDA_Clinical_Diabetes.docx")
